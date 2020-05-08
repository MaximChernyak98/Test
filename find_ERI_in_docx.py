import docx
import os
import openpyxl

#doc = docx.Document('c:\\Work\\Python\\test2\\Протоколы для анализа\\13.03.2020 №224-31.02-04-14.docx')

# Получаем пути файлов для поиска
path_dir = 'c:\\Work\\Python\\test2\\Протоколы для анализа\\'
files = os.listdir(path_dir)
for i in range (0, len(files)):
    files[i] = path_dir + files[i]
print(files)

# позиция для записи в итогово файле
xls_position = 1

# ищет ЭРИ в первой таблице в документе, возвращает номер таблицы и запрашиваемый флюенс
def search_in_first_table (type, docs):
    doc = docx.Document(docs)
    table = doc.tables[1]
    number_of_table = []
    fluence = []
    for row in range(0,len(table.rows)):
        for column in range(0,len(table.columns)):
            cell = (table.cell(row,column)).text
            if cell == type:
                number_of_table.append(int((table.cell(row,0)).text))
                fluence.append((table.cell(row,3)).text)
    if len(fluence) == 0:
        return None, None
    else:
        return number_of_table,fluence
    doc.close()


# ищет дату импульса и полученный флюенс в документе, возвращает дату и полученный флюенс
def get_date_and_fluence (number_table, docs):
    doc = docx.Document(docs)
    table = doc.tables[number_table]
    # ищет последнюю дату
    last_date_text = table.columns[0].cells[-1].text
    last_date = last_date_text.split(',')[-1].strip()
    # ищет флюенс
    for i in range(1, len(table.columns[4].cells)):
        cell_text = table.columns[4].cells[i].text
        if len(str(cell_text)) > 1:
            fluence = cell_text
    return last_date, fluence


# Получает данные поиска и формирует строку для записи в итоговый файл
def return_data (type, doc):
    number_of_table, request_fluence = search_in_first_table (type, doc)
    if not((number_of_table == None) or (request_fluence == None)):
        result = []
        for i in range(0, len(number_of_table)):
            date, actual_fluence = get_date_and_fluence((number_of_table[i] + 1), doc)
            # добавляем 2000 год если нужно
            date_list = date.split(".")
            if len(date_list[2]) < 3:
                date_list[2] = '20' + date_list[2]
            date = '.'.join(date_list)
            result.append(f'{type}, {date}, {request_fluence[i]}, {actual_fluence}, {os.path.split(doc)[1]}')
        return result
    else:
        return None


# Запись строк по одной ЭРИ в итоговый файл
def write_in_xlxs (list_of_data, xls_position):
    from openpyxl import load_workbook
    log = load_workbook('C:\\Work\\Python\\test2\\Результат.xlsx')
    sheet = log.active
    for data in list_of_data:
        a = data.split(", ")
        for i in range (0, len(a)):
            if i == 0 or i == 1:
                sheet.cell(row=xls_position, column=(i+1), value=a[i])
            elif i == 2:
                if 'до' in a[i]:
                    value_list = a[i].replace(',', '.').strip('до ').split('∙10')
                    value = float(value_list[0]) * (10 ** int(value_list[1]))
                    sheet.cell(row=xls_position, column=(i + 1), value=value)
                else:
                    value_list = a[i].replace(',', '.').split('E+')
                    value = float(value_list[0]) * (10 ** int(value_list[1]))
                    sheet.cell(row=xls_position, column=(i + 1), value=value)
                sheet.cell(row=xls_position, column=(i + 1)).number_format = '0.00E+00'
            elif i == 3:
                value_list = a[i].replace(',', '.').split('E+')
                value = float(value_list[0]) * (10 ** int(value_list[1]))
                sheet.cell(row=xls_position, column=(i + 1), value=value)
                sheet.cell(row=xls_position, column=(i + 1)).number_format = '0.00E+00'
            else:
                sheet.cell(row=xls_position, column=(i + 1), value=a[4])
        xls_position += 1
    log.save('C:\\Work\\Python\\test2\\Результат.xlsx')
    log.close()
    return xls_position

# Читает из файла список ЭРИ
def get_list_of_ERI ():
    list_of_ERI = []
    from openpyxl import load_workbook
    log = load_workbook('c:\\Work\\Python\\test2\\Список ЭРИ.xlsx')
    sheet = log.active
    for i in range (1,200):
        element = sheet.cell(row=i, column=1).value
        if not (element == None):
            if not (element in list_of_ERI):
                list_of_ERI.append(element)
    return list_of_ERI


def runner(list, files):
    xls_position = 1
    for element in list:
        for path in files:
            result = return_data(element, path)
            print(result)
            if not (result == None):
                xls_position = write_in_xlxs(result, xls_position)
    return xls_position



list_of_ERI = get_list_of_ERI ()
xls_position = runner (list_of_ERI, files)



